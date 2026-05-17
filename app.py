"""
Smart Document Analyzer - Flask Backend
Features: Upload & Parse, Summarization, Keyword Extraction,
Q&A (Chat with Doc via Gemini), Sentiment Analysis, NER,
Auto-Classification, Table Extraction, Entity Analysis
"""

import os
import re
import math
import collections
from flask import Flask, request, jsonify, render_template
import pdfplumber
from docx import Document
import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np
from dotenv import load_dotenv
from google import genai

# ─────────────────────────────────────────────
# INIT
# ─────────────────────────────────────────────
load_dotenv()

# Line ~26: change this
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
gemini_client = genai.Client(api_key=GEMINI_API_KEY)

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt', 'png', 'jpg', 'jpeg'}

# ─────────────────────────────────────────────
# STOP WORDS
# ─────────────────────────────────────────────
STOP_WORDS = set("""
a about above after again against all also am an and any are aren't as at be because
been before being below between both but by can can't cannot could couldn't did didn't
do does doesn't doing don't down during each few for from further get got had hadn't
has hasn't have haven't having he he'd he'll he's her here here's hers herself him
himself his how how's i i'd i'll i'm i've if in into is isn't it it's its itself
let's me more most mustn't my myself no nor not of off on once only or other ought our
ours ourselves out over own same shan't she she'd she'll she's should shouldn't so
some such than that that's the their theirs them themselves then there there's these
they they'd they'll they're they've this those through to too under until up very was
wasn't we we'd we'll we're we've were weren't what what's when when's where where's
which while who who's whom why why's will with won't would wouldn't you you'd you'll
you're you've your yours yourself yourselves the a also is are was were been being
have has had having do does did doing will would could should may might must shall
can of in on at to for with by about against between into through during before after
above below from up down out off over under again further then once here there when
where why how all both each few more most other some such no nor only own same so
than too very just because if not be
""".lower().split())


# ─────────────────────────────────────────────
# TEXT EXTRACTION
# ─────────────────────────────────────────────
def extract_text_from_pdf(filepath):
    tables = []
    text_pages = []
    with pdfplumber.open(filepath) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            text_pages.append(text)
            page_tables = page.extract_tables()
            for t in page_tables:
                if t and len(t) > 1:
                    tables.append({"page": i + 1, "data": t})
    return "\n".join(text_pages), tables


def extract_text_from_docx(filepath):
    doc = Document(filepath)
    tables = []
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    for table in doc.tables:
        table_data = [[cell.text for cell in row.cells] for row in table.rows]
        tables.append({"page": 1, "data": table_data})
    return "\n".join(full_text), tables


def extract_text_from_txt(filepath):
    with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
        return f.read(), []


def extract_text_from_image(filepath):
    try:
        import pytesseract
        from PIL import Image
        img = Image.open(filepath)
        text = pytesseract.image_to_string(img)
        return text, []
    except Exception as e:
        return f"[OCR not available: {str(e)}]", []


def extract_text(filepath, ext):
    if ext == 'pdf':
        return extract_text_from_pdf(filepath)
    elif ext == 'docx':
        return extract_text_from_docx(filepath)
    elif ext == 'txt':
        return extract_text_from_txt(filepath)
    elif ext in ('png', 'jpg', 'jpeg'):
        return extract_text_from_image(filepath)
    return "", []


# ─────────────────────────────────────────────
# DOCUMENT STRUCTURE DETECTION
# ─────────────────────────────────────────────
def detect_structure(text):
    lines = text.split('\n')
    structure = {
        'headings': [],
        'paragraphs': 0,
        'lists': 0,
        'word_count': 0,
        'char_count': len(text)
    }
    list_pattern = re.compile(r'^\s*[-•*]\s|^\s*\d+[.)]\s', re.MULTILINE)
    for line in lines:
        stripped = line.strip()
        if stripped and re.match(r'^(#{1,6}\s.+|[A-Z][A-Z\s]{4,50}$|\d+\.\s[A-Z])', stripped):
            structure['headings'].append(stripped[:80])
        if stripped and len(stripped) > 40:
            structure['paragraphs'] += 1
    structure['lists'] = len(list_pattern.findall(text))
    structure['word_count'] = len(text.split())
    return structure


# ─────────────────────────────────────────────
# SUMMARIZATION
# ─────────────────────────────────────────────
def tokenize_sentences(text):
    text = re.sub(r'\n+', ' ', text)
    sentences = re.split(r'(?<=[.!?])\s+', text)
    return [s.strip() for s in sentences if len(s.strip()) > 20]


def score_sentences_tfidf(sentences):
    if len(sentences) < 2:
        return np.ones(len(sentences))
    try:
        vectorizer = TfidfVectorizer(stop_words='english', max_features=500)
        tfidf_matrix = vectorizer.fit_transform(sentences)
        scores = np.array(tfidf_matrix.sum(axis=1)).flatten()
        return scores
    except Exception:
        return np.ones(len(sentences))


def summarize(text, mode='short'):
    sentences = tokenize_sentences(text)
    if not sentences:
        return "Could not extract meaningful content."

    scores = score_sentences_tfidf(sentences)
    ranked = sorted(zip(scores, range(len(sentences)), sentences), reverse=True)

    if mode == 'short':
        top_n = min(3, len(sentences))
        top = sorted(ranked[:top_n], key=lambda x: x[1])
        return ' '.join([s for _, _, s in top])

    elif mode == 'bullets':
        top_n = min(7, len(sentences))
        top = sorted(ranked[:top_n], key=lambda x: x[1])
        return [s for _, _, s in top]

    else:  # detailed
        top_n = min(10, max(5, len(sentences) // 5))
        top = sorted(ranked[:top_n], key=lambda x: x[1])
        return ' '.join([s for _, _, s in top])


def eli5_summary(text):
    short = summarize(text, 'short')
    kw = extract_keywords(text, top_n=3)
    topic = kw['keywords'][0] if kw['keywords'] else 'the main topic'
    return (
        f"Here's the simple version: {short}\n\n"
        f"(Think of it like this: the document talks about '{topic}' "
        f"and gives information about it in simple steps.)"
    )


# ─────────────────────────────────────────────
# KEYWORD EXTRACTION
# ─────────────────────────────────────────────
def extract_keywords(text, top_n=15):
    words = re.findall(r'\b[a-zA-Z]{3,}\b', text.lower())
    words = [w for w in words if w not in STOP_WORDS and len(w) > 3]

    freq = collections.Counter(words)
    total = sum(freq.values()) or 1

    sentences = tokenize_sentences(text)
    word_in_docs = collections.Counter()
    for sent in sentences:
        unique = set(re.findall(r'\b[a-zA-Z]{3,}\b', sent.lower()))
        for w in unique:
            word_in_docs[w] += 1

    n_docs = len(sentences) or 1
    tfidf_scores = {}
    for word, count in freq.items():
        tf = count / total
        idf = math.log(n_docs / (word_in_docs.get(word, 1)))
        tfidf_scores[word] = tf * idf

    ranked = sorted(tfidf_scores.items(), key=lambda x: x[1], reverse=True)
    keywords = [w for w, _ in ranked[:top_n]]

    bigrams = []
    for i in range(len(words) - 1):
        if words[i] not in STOP_WORDS and words[i + 1] not in STOP_WORDS:
            bigrams.append(f"{words[i]} {words[i+1]}")
    bigram_freq = collections.Counter(bigrams).most_common(5)

    return {
        "keywords": keywords,
        "bigrams": [b for b, _ in bigram_freq],
        "frequency": dict(freq.most_common(20))
    }


# ─────────────────────────────────────────────
# SENTIMENT ANALYSIS
# ─────────────────────────────────────────────
POSITIVE_WORDS = set("""
good great excellent amazing wonderful fantastic outstanding superb brilliant perfect
positive success successful achieve accomplished effective efficient helpful useful
beneficial valuable important significant impressive notable remarkable love
like enjoy appreciate recommend satisfied happy pleased delighted quality
best better improve improvement gain profit growth strong strengths opportunity
advantage benefit innovative creative solution resolve
""".lower().split())

NEGATIVE_WORDS = set("""
bad poor terrible horrible awful dreadful negative fail failure failed problem issue
concern risk threat weakness difficult hard challenge obstacle barrier limitation
lack missing loss decline decrease reduce error mistake wrong incorrect
inadequate insufficient unsatisfactory disappointing unacceptable harmful dangerous
severe critical urgent serious unfortunately sadly regret complaint ineffective
inefficient costly expensive
""".lower().split())

FORMAL_WORDS = set("therefore thus hence accordingly consequently furthermore moreover additionally pursuant whereby herein thereof".split())
LEGAL_WORDS = set("contract agreement clause section article party parties liability obligation covenant warrant indemnify breach terminate null void enforce".split())
EMOTIONAL_WORDS = set("feel feeling felt emotion love hate fear anger joy sad happy excited worried anxious concerned pleased delighted frustrated angry upset".split())


def analyze_sentiment(text):
    words = re.findall(r'\b[a-zA-Z]+\b', text.lower())
    pos_count = sum(1 for w in words if w in POSITIVE_WORDS)
    neg_count = sum(1 for w in words if w in NEGATIVE_WORDS)
    total = len(words) or 1

    score = (pos_count - neg_count) / total * 100
    if score > 1:
        sentiment = "Positive"
        confidence = min(95, 50 + score * 10)
    elif score < -1:
        sentiment = "Negative"
        confidence = min(95, 50 + abs(score) * 10)
    else:
        sentiment = "Neutral"
        confidence = 65

    formal_count = sum(1 for w in words if w in FORMAL_WORDS)
    legal_count = sum(1 for w in words if w in LEGAL_WORDS)
    emotional_count = sum(1 for w in words if w in EMOTIONAL_WORDS)

    tones = []
    if formal_count > 2:
        tones.append("Formal")
    if legal_count > 3:
        tones.append("Legal")
    if emotional_count > 3:
        tones.append("Emotional")
    if not tones:
        tones.append("Neutral/Informational")

    return {
        "sentiment": sentiment,
        "confidence": round(confidence, 1),
        "positive_score": round(pos_count / total * 100, 2),
        "negative_score": round(neg_count / total * 100, 2),
        "tones": tones,
        "pos_words_found": pos_count,
        "neg_words_found": neg_count
    }


# ─────────────────────────────────────────────
# NAMED ENTITY RECOGNITION
# ─────────────────────────────────────────────
def extract_entities(text):
    entities = {
        "dates": [], "emails": [], "phones": [],
        "urls": [], "organizations": [], "monetary": [],
        "percentages": [], "names": []
    }

    date_patterns = [
        r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b',
        r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{1,2},?\s+\d{4}\b',
        r'\b\d{4}[-/]\d{1,2}[-/]\d{1,2}\b',
        r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b'
    ]
    for p in date_patterns:
        entities["dates"].extend(re.findall(p, text, re.IGNORECASE))

    entities["emails"] = re.findall(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)
    entities["phones"] = re.findall(r'\b(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b', text)
    entities["urls"] = re.findall(r'https?://[^\s]+', text)
    entities["monetary"] = re.findall(
        r'(?:[$€£₹¥]\s*[\d,]+(?:\.\d{2})?|\b[\d,]+(?:\.\d{2})?\s*(?:USD|EUR|GBP|INR|dollars?|euros?|pounds?)\b)',
        text, re.IGNORECASE
    )
    entities["percentages"] = re.findall(r'\b\d+(?:\.\d+)?%\b', text)
    entities["organizations"] = list(set(re.findall(
        r'\b(?:[A-Z][a-z]+ (?:Inc|Corp|Ltd|LLC|Co|Company|Group|Institute|University|College|Association|Organization|Foundation|Department|Ministry|Agency)\b|[A-Z]{2,6}\b)',
        text
    )))[:10]

    names = []
    for sent in text.split('.'):
        found = re.findall(r'(?<!\. )(?<![A-Z])\b([A-Z][a-z]+ [A-Z][a-z]+)\b', sent)
        names.extend(found)
    entities["names"] = list(set(names))[:10]

    for k in entities:
        entities[k] = list(set(entities[k]))[:8]

    return entities


# ─────────────────────────────────────────────
# AUTO CLASSIFICATION
# ─────────────────────────────────────────────
CATEGORY_KEYWORDS = {
    "Resume/CV": ["experience", "skills", "education", "employment", "objective", "curriculum vitae", "references", "bachelor", "master", "degree", "internship", "worked", "position"],
    "Legal Document": ["agreement", "contract", "parties", "clause", "liability", "terms", "obligations", "pursuant", "herein", "whereas", "indemnify", "breach", "arbitration"],
    "Research Paper": ["abstract", "introduction", "methodology", "conclusion", "references", "literature review", "hypothesis", "findings", "journal", "citation", "figure", "table"],
    "Invoice/Financial": ["invoice", "amount", "total", "payment", "due date", "billing", "invoice number", "tax", "subtotal", "quantity", "price", "bill to"],
    "News Article": ["according to", "reported", "sources", "government", "official", "said", "press", "breaking", "today", "yesterday", "news", "journalist"],
    "Technical Document": ["system", "architecture", "implementation", "configuration", "module", "function", "api", "database", "server", "deployment", "technical", "specification"],
    "Medical Document": ["patient", "diagnosis", "treatment", "clinical", "medical", "health", "symptom", "medication", "hospital", "doctor", "disease", "therapy"],
    "Educational Content": ["chapter", "lesson", "exercise", "students", "learning", "course", "curriculum", "assignment", "study", "exam", "quiz", "objective"]
}


def classify_document(text):
    text_lower = text.lower()
    scores = {}
    for category, keywords in CATEGORY_KEYWORDS.items():
        count = sum(1 for kw in keywords if kw in text_lower)
        scores[category] = count / len(keywords)

    sorted_cats = sorted(scores.items(), key=lambda x: x[1], reverse=True)
    primary = sorted_cats[0]
    secondary = sorted_cats[1] if sorted_cats[1][1] > 0.05 else None

    confidence = min(95, int(primary[1] * 300))
    if confidence < 10:
        category = "General Document"
        confidence = 40
    else:
        category = primary[0]

    return {
        "category": category,
        "confidence": confidence,
        "secondary": secondary[0] if secondary else None,
        "all_scores": {k: round(v * 100, 1) for k, v in sorted_cats[:5]}
    }


# ─────────────────────────────────────────────
# Q&A — GEMINI (google-genai SDK)
# ─────────────────────────────────────────────
def answer_question(question, text):
    context = text[:12000] if len(text) > 12000 else text

    prompt = f"""You are a helpful assistant. Answer the user's question based ONLY on the document content provided below.
If the answer is not found in the document, say "I couldn't find that information in the document."

Document:
{context}

Question: {question}

Answer:"""

    try:
        response = gemini_client.models.generate_content(
            model="gemini-2.5-flash",
            contents=prompt
        )
        return {
            "answer": response.text.strip(),
            "confidence": 95,
            "sources": []
        }
    except Exception as e:
        return {
            "answer": f"Gemini error: {str(e)}",
            "confidence": 0,
            "sources": []
        }


# ─────────────────────────────────────────────
# TABLE EXTRACTION
# ─────────────────────────────────────────────
def tables_to_json(tables):
    result = []
    for table_info in tables:
        data = table_info.get("data", [])
        if not data or len(data) < 2:
            continue
        headers = [str(c) if c else "" for c in data[0]]
        rows = []
        for row in data[1:]:
            clean_row = [str(c) if c else "" for c in row]
            if any(c.strip() for c in clean_row):
                rows.append(clean_row)
        result.append({
            "page": table_info.get("page", 1),
            "headers": headers,
            "rows": rows[:50]
        })
    return result


# ─────────────────────────────────────────────
# ROUTES
# ─────────────────────────────────────────────
@app.route('/')
def index():
    return render_template('index.html')


@app.route('/analyze', methods=['POST'])
def analyze():
    if 'file' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No file selected"}), 400

    ext = file.filename.rsplit('.', 1)[-1].lower() if '.' in file.filename else ''
    if ext not in ALLOWED_EXTENSIONS:
        return jsonify({"error": f"Unsupported file type: {ext}"}), 400

    filepath = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
    file.save(filepath)

    try:
        text, raw_tables = extract_text(filepath, ext)
        if not text or len(text.strip()) < 50:
            return jsonify({"error": "Could not extract meaningful text from document."}), 400

        structure = detect_structure(text)
        reading_time = max(1, round(structure['word_count'] / 200))

        return jsonify({
            "success": True,
            "filename": file.filename,
            "text_preview": text[:1000],
            "full_text": text,
            "structure": structure,
            "reading_time": reading_time,
            "summary": {
                "short": summarize(text, 'short'),
                "detailed": summarize(text, 'detailed'),
                "bullets": summarize(text, 'bullets'),
                "eli5": eli5_summary(text)
            },
            "keywords": extract_keywords(text),
            "sentiment": analyze_sentiment(text),
            "entities": extract_entities(text),
            "classification": classify_document(text),
            "tables": tables_to_json(raw_tables)
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        if os.path.exists(filepath):
            os.remove(filepath)


@app.route('/chat', methods=['POST'])
def chat():
    data = request.get_json()
    question = data.get('question', '').strip()
    text = data.get('text', '').strip()

    if not question or not text:
        return jsonify({"error": "Missing question or document text"}), 400

    result = answer_question(question, text)
    return jsonify(result)


@app.route('/export_table', methods=['POST'])
def export_table():
    data = request.get_json()
    table = data.get('table')
    fmt = data.get('format', 'csv')

    if not table:
        return jsonify({"error": "No table data"}), 400

    headers = table.get('headers', [])
    rows = table.get('rows', [])
    df = pd.DataFrame(rows, columns=headers if headers else None)

    if fmt == 'csv':
        return jsonify({"data": df.to_csv(index=False), "format": "csv"})
    else:
        return jsonify({"data": df.to_json(orient='records'), "format": "json"})


if __name__ == '__main__':
    app.run(debug=True, port=5050)